import streamlit as st
from docx import Document
import tempfile
import zipfile
import io

from extractor import extract_document, detect_sections, extract_section

st.set_page_config(page_title="AI PDF Content Extractor", layout="wide")

st.title("AI PDF Content Extractor")
st.divider()

st.write("Upload a PDF, select one or more main sections from the Table of Contents, preview the extracted content, and download the results.")

uploaded_file = st.file_uploader("Upload PDF file", type=["pdf"])

if uploaded_file:

    with st.spinner("Processing PDF..."):
        doc = extract_document(uploaded_file)

    sections = detect_sections(doc)

    if sections:
        st.subheader("Select Sections")

        selected_sections = []

        for i, sec in enumerate(sections):
            label = sec["title"]
            if st.checkbox(label, key=f"sec_{i}_{sec['num']}"):
                selected_sections.append(sec)

        if st.button("🚀 Extract Selected Sections"):

            zip_buffer = io.BytesIO()
            combined_text = ""

            st.subheader("Preview")

            with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zipf:

                for i, sec in enumerate(selected_sections):
                    content = extract_section(doc, sec)

                    st.markdown(f"### {sec['title']}")
                    st.text_area(
                        label=f"Content {i}",
                        value=content,
                        height=260,
                        key=f"preview_{i}_{sec['num']}"
                    )

                    combined_text += f"{sec['title']}\n\n{content}\n\n"

                # Create Word output
                word_doc = Document()
                word_doc.add_heading("Extracted PDF Sections", level=1)
                word_doc.add_paragraph(combined_text)

                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                word_doc.save(tmp.name)

                zipf.write(tmp.name, "extracted_sections.docx")

            st.success("Extraction completed.")

            st.download_button(
                label="⬇ Download Word File",
                data=zip_buffer.getvalue(),
                file_name="sections.zip",
                mime="application/zip"
            )

    else:
        st.warning("No valid Table of Contents sections were detected in this PDF.")