import streamlit as st
from docx import Document
import tempfile
import zipfile
import io

from extractor import (
    extract_document,
    detect_sections,
    locate_section_starts,
    extract_section,
    clean_for_word,
)

st.set_page_config(page_title="AI PDF Content Extractor", layout="wide")

st.title("AI PDF Content Extractor")
st.divider()

st.write(
    "Upload one or more PDFs, select main sections from the Table of Contents, "
    "preview extracted content, and download the results."
)

uploaded_files = st.file_uploader(
    "Upload PDF file(s)",
    type=["pdf"],
    accept_multiple_files=True
)

if uploaded_files:

    all_results = []

    for file_index, uploaded_file in enumerate(uploaded_files):
        st.markdown("---")
        st.subheader(f"📄 {uploaded_file.name}")

        with st.spinner(f"Processing {uploaded_file.name}..."):
            doc = extract_document(uploaded_file)
            toc_sections = detect_sections(doc)
            located_sections = locate_section_starts(doc, toc_sections)

        if located_sections:
            st.markdown("**Select Sections**")

            selected_sections = []

            for sec_index, sec in enumerate(located_sections):
                label = sec["title"]
                if st.checkbox(
                    label,
                    key=f"sec_{file_index}_{sec_index}_{sec['num']}"
                ):
                    selected_sections.append(sec)

            all_results.append({
                "file_name": uploaded_file.name,
                "doc": doc,
                "located_sections": located_sections,
                "selected_sections": selected_sections
            })
        else:
            st.warning(
                f"No valid Table of Contents sections were detected (or mapped to body headings) in {uploaded_file.name}."
            )

    if st.button("🚀 Extract Selected Sections From All PDFs"):

        zip_buffer = io.BytesIO()

        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED) as zipf:

            for result_index, result in enumerate(all_results):
                file_name = result["file_name"]
                doc = result["doc"]
                located_sections = result["located_sections"]
                selected_sections = result["selected_sections"]

                if not selected_sections:
                    continue

                st.markdown("---")
                st.subheader(f"Preview — {file_name}")

                # Create one Word file per PDF
                word_doc = Document()
                word_doc.add_heading(f"Extracted PDF Sections - {file_name}", level=1)

                for sec_index, sec in enumerate(selected_sections):
                    content = extract_section(doc, located_sections, sec)
                    safe_content = clean_for_word(content)

                    st.markdown(f"### {sec['title']} ({file_name})")
                    st.text_area(
                        label=f"Content_{result_index}_{sec_index}",
                        value=safe_content,
                        height=280,
                        key=f"preview_{result_index}_{sec_index}_{sec['num']}"
                    )

                    word_doc.add_heading(sec["title"], level=2)

                    if safe_content:
                        for para in safe_content.split("\n\n"):
                            para = clean_for_word(para)
                            if para.strip():
                                word_doc.add_paragraph(para)
                    else:
                        word_doc.add_paragraph("[No content extracted]")

                tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
                word_doc.save(tmp.name)

                output_name = file_name.rsplit(".", 1)[0] + "_sections.docx"
                zipf.write(tmp.name, output_name)

        st.success("Extraction completed.")

        st.download_button(
            label="⬇ Download All Word Files (ZIP)",
            data=zip_buffer.getvalue(),
            file_name="all_extracted_sections.zip",
            mime="application/zip"
        )